"""File processing utilities for extracting text from various formats"""
import io
from typing import Tuple
import PyPDF2
from docx import Document
import csv


async def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        text = ""
        for page in pdf_reader.pages:
            page_text = page.extract_text() or ""
            text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error reading PDF: {str(e)}")


async def extract_text_from_csv(file_content: bytes) -> str:
    """Extract text from CSV file - includes headers and sample rows"""
    try:
        text_content = file_content.decode('utf-8', errors='ignore')
        lines = text_content.split('\n')
        
        # Use first 50 lines or all lines if less
        sample_lines = lines[:50]
        extracted = '\n'.join(sample_lines)
        
        return extracted.strip()
    except Exception as e:
        raise ValueError(f"Error reading CSV: {str(e)}")


async def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = Document(io.BytesIO(file_content))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " | "
                text += "\n"
        return text.strip()
    except Exception as e:
        raise ValueError(f"Error reading DOCX: {str(e)}")


async def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file"""
    try:
        return file_content.decode('utf-8', errors='ignore').strip()
    except Exception as e:
        raise ValueError(f"Error reading TXT: {str(e)}")


async def process_uploaded_file(file_content: bytes, filename: str) -> str:
    """Process uploaded file and extract text based on file type"""
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.pdf'):
        text = await extract_text_from_pdf(file_content)
    elif filename_lower.endswith('.csv'):
        text = await extract_text_from_csv(file_content)
    elif filename_lower.endswith('.docx'):
        text = await extract_text_from_docx(file_content)
    elif filename_lower.endswith('.txt'):
        text = await extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Unsupported file type: {filename}")
    
    # Truncate to ~2000 chars for token efficiency
    if len(text) > 2000:
        text = text[:2000] + "...[truncated]"
    
    return text
